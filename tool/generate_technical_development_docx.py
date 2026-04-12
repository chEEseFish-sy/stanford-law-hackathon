from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = "/Users/curtis/Desktop/hackson/product-design/VeriCap-技术开发流程文档.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "D9EAF7")
        if widths and i < len(widths):
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if widths and i < len(widths):
                cells[i].width = widths[i]
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Arial"
styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Title"].font.size = Pt(22)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Heading 1"].font.size = Pt(15)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Heading 2"].font.size = Pt(12.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.style = "Title"
title.add_run("VeriCap\n技术开发流程文档")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("版本：Web 应用技术实现方案与模块拆解\n")
subtitle.add_run("日期：2026-04-12")

doc.add_paragraph("")

doc.add_heading("1. 文档目标", level=1)
add_bullets(
    doc,
    [
        "本文档用于定义 VeriCap 的 Web 应用技术开发流程，明确技术栈、系统分层、模块职责、开发阶段与验收重点。",
        "该文档服务对象包括产品负责人、UX Agent、开发 Agent、后端工程师、前端工程师以及后续联调与验证成员。",
        "重点补充一个可回溯的拓扑数据结构方案，用于支撑 working cap table 的版本化生成、差异解释与状态回溯。",
    ],
)

doc.add_heading("2. 技术方案总览", level=1)
add_bullets(
    doc,
    [
        "VeriCap 推荐采用“Web Workbench + AI/规则混合解析 + 事件重建引擎 + 版本回溯图”的架构。",
        "整体系统目标不是一次性输出法律终局结论，而是持续生成可追溯、可解释、可回放的 working cap table。",
        "系统核心主线应保持为：文件进入 → 文件整理 → 证据审理 → 事件重建 → cap table 生成 → 对账异常 → 版本回溯 → 报告导出。",
        "技术实现上应优先保障数据结构清晰、状态可回放、异步任务稳定、前端工作台可解释，而不是过早追求复杂证券全覆盖。",
    ],
)

doc.add_heading("3. 推荐技术栈", level=1)
add_table(
    doc,
    ["层级", "推荐技术", "作用", "选择原因"],
    [
        ["前端框架", "Next.js + TypeScript", "搭建 Web 工作台、路由、SSR/CSR 混合渲染", "适合复杂工作台型产品，开发效率高，生态成熟"],
        ["UI 层", "Tailwind CSS + shadcn/ui", "统一样式系统、表格面板、表单、弹层与组件复用", "适合企业型工作台，定制效率高"],
        ["前端状态", "Zustand", "管理工作台本地 UI 状态、当前选中文档、侧边解释面板、筛选条件", "轻量、清晰，适合局部复杂状态"],
        ["服务端数据获取", "TanStack Query", "管理 API 请求、缓存、轮询和任务状态刷新", "适合异步文档处理与多面板数据同步"],
        ["后端 API", "FastAPI + Pydantic", "提供文档上传、工作区管理、审理结果、回溯与导出接口", "适合 AI、规则引擎、计算服务集成"],
        ["异步任务", "Celery + Redis", "处理 OCR、解析、抽取、事件重建、导出等后台任务", "将耗时流程与交互解耦，便于任务排队和重试"],
        ["主数据库", "PostgreSQL", "存储 workspace、文件元数据、证据、事件、版本、审计日志", "关系清晰、事务能力强，适合版本和图关系持久化"],
        ["向量检索", "pgvector", "支持条款语义召回、解释引用和相似证据定位", "在 PostgreSQL 内扩展，降低系统复杂度"],
        ["对象存储", "S3 兼容对象存储 / 腾讯云 COS", "存储原始文件、预处理文件、导出结果", "适合大文件和版本留存"],
        ["文档解析", "Python 文档解析组件 + OCR 服务适配层", "读取 PDF、Word、Excel，生成可消费中间结构", "便于与 AI 解析流程结合"],
        ["鉴权", "JWT + HttpOnly Cookie + RBAC", "登录、权限隔离、workspace 访问控制", "足够满足首版安全需求，便于多角色扩展"],
        ["部署", "Vercel 前端 + 容器化后端 + 托管 PostgreSQL", "形成可快速迭代的 Web 发布链路", "适合早期试点与快速上线"],
        ["监控", "Sentry + OpenTelemetry + 结构化日志", "错误追踪、任务追踪、链路诊断", "保证文档处理链路可观测"],
    ],
    widths=[Inches(1.15), Inches(1.75), Inches(2.0), Inches(2.2)],
)

doc.add_heading("4. 系统分层架构", level=1)
add_numbered(
    doc,
    [
        "表现层：Web 工作台，负责 workspace、文件、冲突、event timeline、working cap table、change review 与导出界面。",
        "应用层：API/BFF，负责聚合权限、任务状态、查询视图和前端需要的结构化响应。",
        "领域层：文件整理、证据审理、事件重建、回放与对账等核心业务服务。",
        "任务层：OCR、抽取、解析、图计算、导出等异步处理流水线。",
        "存储层：PostgreSQL 保存结构化事实与回溯图，对象存储保存原始与派生文件。",
        "智能层：LLM / 规则引擎 / 检索增强，为字段抽取、解释生成和冲突建议提供能力。",
    ],
)

doc.add_heading("5. 回溯拓扑数据结构设计", level=1)
doc.add_paragraph(
    "VeriCap 的状态回溯不建议只用“线性版本号 + 全量快照”实现，而应采用“Version Graph + Event DAG”的拓扑结构。该结构能同时支持事件依赖、人工决策、版本分支、状态重放与差异解释。"
)

doc.add_heading("5.1 设计目标", level=2)
add_bullets(
    doc,
    [
        "支持用户回到任意历史审理状态，重新生成 working cap table。",
        "支持新证据进入后从某个历史节点分叉，形成新的版本分支。",
        "支持解释某一行 cap table 是由哪些事件、哪些证据、哪些用户决策共同产生。",
        "支持比较两个版本之间发生了哪些变化，以及变化是由哪些节点引起。",
    ],
)

doc.add_heading("5.2 核心节点类型", level=2)
add_table(
    doc,
    ["节点类型", "说明", "核心字段", "作用"],
    [
        ["DocumentNode", "原始文档或某个文档版本", "document_id、file_type、uploaded_at、supersedes", "承载原始证据源"],
        ["EvidenceNode", "从文档中抽出的可审理证据单元", "evidence_id、field_name、value、confidence、source_range", "承接字段级冲突与审理"],
        ["DecisionNode", "人工决策记录", "decision_id、action、actor、reason、created_at", "记录接受、排除、暂缓等动作"],
        ["EventNode", "标准化股权事件", "event_id、event_type、event_date、inputs", "作为 cap table 计算的业务事实单元"],
        ["SnapshotNode", "某次 working cap table 结果快照", "snapshot_id、version_no、hash、generated_at", "用于查看与回放特定版本结果"],
        ["VersionNode", "一次完整状态版本入口", "version_id、base_version_id、status、created_at", "管理版本分支与回溯入口"],
    ],
    widths=[Inches(1.2), Inches(2.1), Inches(2.2), Inches(1.6)],
)

doc.add_heading("5.3 核心边类型", level=2)
add_table(
    doc,
    ["边类型", "起点 → 终点", "含义"],
    [
        ["derived_from", "EvidenceNode → DocumentNode", "该证据来自哪份文档或哪一段内容"],
        ["resolved_by", "EvidenceNode → DecisionNode", "该证据冲突由哪个人工决策处理"],
        ["supports", "EventNode → EvidenceNode", "某个事件由哪些证据支持"],
        ["depends_on", "EventNode → EventNode", "某个事件依赖前置事件才能生效"],
        ["produces", "SnapshotNode → EventNode", "该快照由哪些已生效事件生成"],
        ["branches_from", "VersionNode → VersionNode", "该版本从哪个历史版本分叉而来"],
        ["points_to", "VersionNode → SnapshotNode", "该版本当前展示的结果快照是什么"],
    ],
    widths=[Inches(1.3), Inches(2.3), Inches(2.8)],
)

doc.add_heading("5.4 回放与回溯机制", level=2)
add_numbered(
    doc,
    [
        "用户进入某个 VersionNode。",
        "系统读取该版本下所有已生效 DecisionNode。",
        "系统筛选出被接受的 EvidenceNode 与可成立的 EventNode。",
        "系统对 EventNode 按“业务日期 + 依赖关系 + 决策状态”执行拓扑排序。",
        "按排序结果依次重放事件，生成新的 SnapshotNode。",
        "若用户回退到历史版本并修改某个决策，则从该节点重新分叉，生成新的 VersionNode 与 SnapshotNode。",
    ],
)

doc.add_heading("5.5 为什么采用拓扑结构", level=2)
add_bullets(
    doc,
    [
        "cap table 的生成本质上不是纯线性流水，而是“带依赖关系的事件重放”。",
        "同一份新证据可能只影响部分事件和下游结果，使用图结构可以避免全量重新理解所有历史。",
        "当用户需要解释某一项结果时，图结构可天然返回上游事件、证据和决策链路。",
        "当用户从某个历史状态重新判断时，图分支比线性版本号更适合表达“回到过去并产生新未来”的场景。",
    ],
)

doc.add_heading("6. 技术开发流程", level=1)
add_table(
    doc,
    ["阶段", "目标", "关键输出", "验收重点"],
    [
        ["Phase 0 领域建模", "锁定状态模型、事件模型和版本回溯图", "领域对象、ER 图、Version Graph 设计", "核心实体与边界无冲突"],
        ["Phase 1 基础工程", "搭建前后端骨架与基础环境", "前端工程、后端工程、数据库迁移、对象存储接入", "开发环境可跑通"],
        ["Phase 2 文件进入链路", "完成上传、存储、预处理、任务编排", "上传接口、文件状态机、异步任务流水线", "文件可进入可追踪处理状态"],
        ["Phase 3 证据抽取与审理", "把文档转成证据单元并可人工审理", "EvidenceNode、审理队列、冲突视图", "证据状态可被记录和回写"],
        ["Phase 4 事件重建引擎", "从证据生成事件图", "EventNode、依赖关系、拓扑排序器", "事件可稳定重放"],
        ["Phase 5 cap table 生成与对账", "跑通核心业务闭环", "working cap table、reported 对账、异常类型", "最小闭环可演示"],
        ["Phase 6 回溯与解释", "支持版本回放、差异解释和 change review", "VersionNode、SnapshotNode、解释接口", "用户能回到历史状态并理解变化原因"],
        ["Phase 7 导出与运维", "完善报告导出、权限、监控与审计", "导出链路、操作日志、监控告警", "可试点、可追踪、可维护"],
    ],
    widths=[Inches(1.2), Inches(1.75), Inches(1.9), Inches(1.65)],
)

doc.add_heading("7. 模块说明", level=1)

modules = [
    (
        "模块 1：Auth & Workspace",
        [
            "完成用户登录、角色权限、workspace 创建与成员隔离。",
            "管理 matter 基础信息、当前处理状态、最近操作记录。",
            "为整个工作台提供统一入口和权限边界。",
        ],
    ),
    (
        "模块 2：Document Intake",
        [
            "完成 PDF、Word、Excel 与压缩文件包上传。",
            "负责文件命名标准化、对象存储写入、状态初始化和处理任务入队。",
            "输出 DocumentNode 与基础文件元数据。",
        ],
    ),
    (
        "模块 3：Document Parsing & OCR",
        [
            "完成扫描件 OCR、文本抽取、页面切分、基础结构识别。",
            "输出面向抽取服务可消费的中间格式。",
            "记录解析质量、异常与可回溯的原始片段位置。",
        ],
    ),
    (
        "模块 4：Evidence Extraction",
        [
            "从文档中抽取字段级证据单元，生成 EvidenceNode。",
            "标注字段名、值、置信度、来源页码与文本范围。",
            "为后续冲突检测和事件重建提供标准化输入。",
        ],
    ),
    (
        "模块 5：Evidence Review",
        [
            "构建冲突队列、缺失项提示和人工审理动作。",
            "支持 Confirmed、Needs Review、Excluded 三种状态。",
            "将用户动作沉淀为 DecisionNode，作为下游计算输入。",
        ],
    ),
    (
        "模块 6：Event Reconstruction Engine",
        [
            "基于已接受证据生成 EventNode。",
            "维护事件依赖关系、适用条件、前提假设与业务日期。",
            "提供拓扑排序和事件重放所需的核心能力。",
        ],
    ),
    (
        "模块 7：Version Graph & Replay Engine",
        [
            "维护 VersionNode、SnapshotNode 与分支关系。",
            "支持从任意历史节点重新回放、生成新分支与差异结果。",
            "这是 VeriCap 支撑状态回溯和 cap table 生成回溯的核心技术模块。",
        ],
    ),
    (
        "模块 8：Cap Table Generator",
        [
            "按事件拓扑顺序重放，生成 working cap table。",
            "支持 current / fully diluted 双视图与结果状态分层。",
            "输出行级来源映射，供解释和审计使用。",
        ],
    ),
    (
        "模块 9：Reconciliation & Diff",
        [
            "将 working cap table 与 reported cap table 做系统比对。",
            "识别 count mismatch、missing conversion、missing approval、timeline inconsistency 等异常。",
            "输出差异摘要、影响范围和待处理项。",
        ],
    ),
    (
        "模块 10：Explanation & Audit Report",
        [
            "提供 explain this row、异常解释、版本差异解释。",
            "生成问题清单、证据索引和导出报告。",
            "确保所有重要结果具备 evidence-backed 的解释链路。",
        ],
    ),
    (
        "模块 11：Frontend Workbench",
        [
            "实现 Workspace Dashboard、Document Organization、Evidence Review Queue、Event Timeline、Working Cap Table、Change Review、Export。",
            "管理筛选、侧边栏、回流、轮询与关键状态页。",
            "保证复杂信息以工作台方式呈现，而不是碎片页面堆叠。",
        ],
    ),
    (
        "模块 12：Observability & Audit Trail",
        [
            "记录任务链路、失败原因、操作日志和版本变更日志。",
            "提供错误追踪、性能监控和任务重试观察能力。",
            "保证文档处理链路可诊断、可恢复、可问责。",
        ],
    ),
]

for heading, bullets in modules:
    doc.add_heading(heading, level=2)
    add_bullets(doc, bullets)

doc.add_heading("8. 模块依赖顺序建议", level=1)
add_numbered(
    doc,
    [
        "先完成 Auth & Workspace、Document Intake、基础数据库迁移。",
        "再完成 Document Parsing、Evidence Extraction、Evidence Review。",
        "随后实现 Event Reconstruction Engine 与 Version Graph & Replay Engine。",
        "在此基础上完成 Cap Table Generator 与 Reconciliation & Diff。",
        "最后补齐 Explanation、Export、Observability 与更完整的前端工作台细节。",
    ],
)

doc.add_heading("9. 数据库与核心表建议", level=1)
add_table(
    doc,
    ["表/集合", "核心内容", "说明"],
    [
        ["workspaces", "workspace 基础信息、状态、拥有者", "matter 级隔离入口"],
        ["documents", "文档元数据、对象存储地址、版本关系", "对应 DocumentNode"],
        ["evidences", "字段证据、来源区间、置信度、状态", "对应 EvidenceNode"],
        ["decisions", "审理动作、原因、执行人", "对应 DecisionNode"],
        ["events", "标准化股权事件、日期、依赖、状态", "对应 EventNode"],
        ["versions", "版本入口、父版本、状态", "对应 VersionNode"],
        ["snapshots", "working cap table 快照、摘要、哈希", "对应 SnapshotNode"],
        ["reported_cap_tables", "用户提供的 cap table 数据", "用于对账"],
        ["audit_logs", "操作与系统日志", "用于追踪和审计"],
        ["jobs", "异步任务状态、重试次数、失败原因", "用于任务可观测性"],
    ],
    widths=[Inches(1.45), Inches(2.2), Inches(2.25)],
)

doc.add_heading("10. API 与前后端交互原则", level=1)
add_bullets(
    doc,
    [
        "长耗时动作必须异步化，前端通过任务状态轮询或订阅刷新结果。",
        "所有关键查询接口都应返回 summary + detail 的双层数据，避免前端首次加载过重。",
        "解释类接口必须支持按 row、按 event、按 version 三种维度请求。",
        "版本回溯接口必须支持：查看某版本、从某版本创建分支、比较两个版本。",
        "前端不应自己重建业务真相，只消费后端提供的结构化视图对象。",
    ],
)

doc.add_heading("11. 开发阶段验收标准", level=1)
add_bullets(
    doc,
    [
        "文件上传后能看到明确状态，并能追踪处理进度。",
        "系统能生成 Evidence Review Queue，并允许人工决策回写。",
        "系统能基于证据生成事件，并通过拓扑排序重放。",
        "系统能生成 working cap table，并发现至少三类核心异常。",
        "系统能从历史版本重新回放，生成新的分支结果。",
        "系统能解释某一项结果来自哪些证据、哪些事件和哪些用户决策。",
        "系统能导出结构化报告，并保留完整审计链路。",
    ],
)

doc.add_heading("12. 风险与技术注意事项", level=1)
add_bullets(
    doc,
    [
        "不要把版本回溯实现成仅依赖前端状态的临时逻辑，必须落库并可追踪。",
        "不要把事件重放写死在单一线性流程里，必须允许依赖关系和版本分叉。",
        "不要让前端自行拼接证据解释链，解释结果必须由后端统一生成。",
        "大文件处理、OCR、导出等操作必须异步化，否则会严重影响交互体验。",
        "所有 AI 抽取结果都要保留来源区间和置信度，不能只保留最终字段值。",
    ],
)

doc.add_heading("13. Handoff Summary", level=1)
add_bullets(
    doc,
    [
        "给产品与 UX：VeriCap 的核心不是大表展示，而是风险优先级、证据追溯、事件解释和版本回溯。",
        "给开发：Version Graph + Event DAG 是本项目的核心技术底座，必须优先设计清楚。",
        "给前端：工作台页面必须直接消费结构化视图，不要在浏览器侧重建业务真相。",
        "给后端：文档处理、证据审理、事件重建、回放引擎和解释服务是最优先的核心模块。",
        "给整个团队：首版目标是跑通高置信度、可解释、可回溯的 working cap table 闭环，而不是覆盖所有复杂证券规则。",
    ],
)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("VeriCap 技术开发流程文档")

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
