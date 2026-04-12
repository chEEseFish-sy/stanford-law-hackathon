from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


output_path = "/Users/curtis/Desktop/hackson/Cap_Table_Audit_Copilot_市场调研报告.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "D9EAF7")
        if widths and i < len(widths):
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if widths and i < len(widths):
                cells[i].width = widths[i]
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Arial"
styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Title"].font.size = Pt(22)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Heading 1"].font.size = Pt(15)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
styles["Heading 2"].font.size = Pt(12.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.style = "Title"
title.add_run("Cap Table Audit Copilot\n市场调研报告")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("版本：基于 Idea 1 / Cap Table Audit Copilot 的商业化与竞品研究\n")
subtitle.add_run("日期：2026-04-12")

doc.add_paragraph("")

doc.add_heading("1. 执行摘要", level=1)
add_bullets(
    doc,
    [
        "Cap Table Audit Copilot 处在“法律科技 + 融资尽调 + 股权管理”交叉带，价值主张不是泛化文档问答，而是基于法律文件重建理论 cap table，并对现有 cap table 做可追溯的一致性审计。",
        "从外部市场看，LegalTech 仍处于较快增长期。Grand View Research 显示，全球 LegalTech 市场 2025 年约为 287 亿美元，预计 2033 年达到约 697 亿美元，2026-2033 年 CAGR 为 12.2%，其中律所仍是最大终端客户群。",
        "从需求成熟度看，AI 在合同审阅、条款抽取、尽调辅助等场景已被证明能显著提升效率。美国律师协会 2025 年文章汇总的研究显示，在部分结构化合同任务中，AI 已达到或超过人工表现，同时显著缩短审阅时间。",
        "当前市场已有强势平台覆盖 cap table 管理、409A、股权激励、合同审阅和尽调，但“以法律文件为证据源的股权事件重建 + 异常审计”仍缺乏非常明确的标准化产品领导者。",
        "因此，该产品的最优切入方式不是与通用 CLM、通用法律助手正面对打，而是锁定 Delaware C-Corp 早期融资、VC/创业律师和 in-house legal/finance ops 的“高价值、低容错、强证据链”场景。",
    ],
)

doc.add_heading("2. 研究对象与方法", level=1)
add_bullets(
    doc,
    [
        "研究对象：AI-powered cap table reconciliation and equity audit copilot for startup financings。",
        "输入材料：/Users/curtis/Desktop/hackson/idea_1_cap_table.md。",
        "研究方法：对产品定义、目标用户、MVP 范围进行抽象；结合 LegalTech 行业报告、美国律师协会公开文章及 Carta、Pulley、Ledgy、Shareworks、Litera/Kira、Eqvista 等公开资料完成竞品与市场环境分析。",
        "边界说明：由于“cap table 审计 copilot”是交叉型新类别，公开市场很少给出完全同口径的细分市场规模；因此报告在市场规模部分采用“行业数据 + 邻近赛道 + bottom-up 假设”的方式进行估算，并明确保留假设条件。",
    ],
)

doc.add_heading("3. 产品与市场定义", level=1)
doc.add_paragraph(
    "Cap Table Audit Copilot 的核心不是做单文档摘要，而是把 SAFE、SPA、Option Grant、Board Consent、Convertible Note 与公司 cap table 统一抽象为股权事件，再按时间顺序重放，输出理论 cap table、冲突列表与证据链。"
)
add_bullets(
    doc,
    [
        "类别归属：LegalTech / Transaction Intelligence / Equity Audit / Diligence Workflow。",
        "主要价值：降低人工 cross-check 成本，提前暴露高风险不一致，提升材料可信度与融资准备效率。",
        "购买动因：融资尽调压力、版本控制混乱、SAFE/note 转换复杂、审批链缺失风险、审计追溯要求。",
        "理想切口：Delaware C-Corp 早期融资文件包；由律所融资律师、paralegal、in-house legal/finance ops 高频使用。",
    ],
)

doc.add_heading("4. 宏观市场机会", level=1)
add_bullets(
    doc,
    [
        "LegalTech 增长仍然稳健。Grand View Research 公开数据显示，全球 LegalTech 市场 2025 年约为 287.4 亿美元，2033 年预计约 696.9 亿美元，2026-2033 年 CAGR 为 12.2%。",
        "律所端仍是最大付费主体。该报告显示，law firms 在 2025 年占据最大终端市场份额，说明以律师工作流为中心切入仍是商业化上最容易被理解的路径。",
        "AI 与分析能力的重要性上升。Grand View Research 指出 analytics 是增长最快的细分方向之一，CLM 也是当前收入规模较大的类别，说明“自动抽取 + 分析 + 风险洞察”已经是市场接受的价值组合。",
        "客户对速度、透明度与可审计性要求上升。美国律师协会 2025 年文章指出，AI 的价值已从单点生成延伸到分析、条款抽取、工作流集成与标准化控制。",
    ],
)

doc.add_heading("5. 用户画像与需求强度", level=1)
add_table(
    doc,
    ["用户群", "核心任务", "痛点强度", "付费意愿判断", "适配度"],
    [
        [
            "融资律师 / Paralegal",
            "快速核对交易文件与 cap table 是否一致",
            "极高：文件多、时间紧、责任重",
            "高：节省 billable hours，提升首轮 review 吞吐",
            "最高",
        ],
        [
            "In-house legal / finance ops",
            "在外部尽调前做内部自检",
            "高：资料分散、历史遗留问题多",
            "中高：可减少交易摩擦与补件成本",
            "高",
        ],
        [
            "Founder / CFO",
            "融资前验证股权表可信度",
            "中高：缺专业能力但风险感知强",
            "中：更关注性价比与易用性",
            "中高",
        ],
        [
            "VC / PE 投后或法务支持团队",
            "快速理解被投企业历史股权结构",
            "中：需求存在但不是所有团队高频",
            "中",
            "中",
        ],
    ],
    widths=[Inches(1.4), Inches(2.15), Inches(1.5), Inches(1.5), Inches(0.8)],
)

doc.add_paragraph("")
add_bullets(
    doc,
    [
        "Primary persona 应继续锁定：startup financing lawyers and legal support teams。",
        "这类用户最重视的不是“AI 写得像不像律师”，而是“能不能快速找出问题、证据在哪、我还需要补什么”。",
        "购买与扩张顺序建议：先打律所与高成长公司法务/财务，再扩张到 founder 自查、投后管理与审计协作。",
    ],
)

doc.add_heading("6. 需求驱动因素", level=1)
add_numbered(
    doc,
    [
        "股权结构越来越复杂：SAFE、convertible note、option pool、board consent、multiple closing 等叠加，传统 Excel 难以维持真实单一版本。",
        "融资与尽调节奏加快：外部律师、公司管理层、投资方都希望更快形成第一轮风险判断。",
        "证据可追溯要求提升：结论必须落回具体文件、页码、条款，而不是只有抽象摘要。",
        "AI 审阅能力成熟：条款抽取、对比和总结已有较强可用性，使“从文档到事件模型”的工程化路径更可行。",
        "客户对风险控制更敏感：任何 cap table 错误都可能影响估值、交易节奏、员工股权沟通与投资人信任。",
    ],
)

doc.add_heading("7. 竞品格局", level=1)
doc.add_paragraph(
    "当前市场并不存在一个与 Cap Table Audit Copilot 完全同构的绝对龙头；真实竞争来自三类产品：股权管理平台、法律尽调/合同审阅平台、以及通用 AI 法务工作台。"
)
add_table(
    doc,
    ["竞品", "类别", "核心强项", "与本产品的重叠", "关键空白 / 机会"],
    [
        [
            "Carta",
            "股权管理平台",
            "cap table、409A、融资建模、合规、生态位强",
            "覆盖股权数据管理与部分审计准备",
            "更偏“记录与运营系统”，不是从外部法律文件重建股权历史",
        ],
        [
            "Pulley",
            "创业公司股权管理",
            "早期公司友好、定价相对透明、SAFE/409A/建模强",
            "覆盖 founder/CFO 的 equity workflow",
            "仍更偏管理工具，不以法律证据链和异常审计为主卖点",
        ],
        [
            "Ledgy",
            "全球股权与合规管理",
            "单一数据源、合规、动态交易与文档关联、员工股权体验强",
            "与 cap table 清洁度、历史交易可视化有重合",
            "更偏股权计划与合规运营，对融资法律文件审计不是核心叙事",
        ],
        [
            "Shareworks",
            "企业级股权管理",
            "大型企业/跨地区、报告、409A、实时数据源",
            "覆盖成熟企业 cap table 与 reporting",
            "企业级厚平台，未聚焦“创业融资文件 → 重建 → 异常审计”的律师场景",
        ],
        [
            "Litera Kira",
            "法律尽调 / 合同审阅 AI",
            "条款抽取、 due diligence、成熟律所渗透",
            "在文档抽取和尽调工作流高度相关",
            "擅长合同层面的 review，不直接输出理论 cap table 与股权一致性审计",
        ],
        [
            "Eqvista",
            "股权管理 + 409A",
            "性价比、409A、cap table、文档管理",
            "服务中小公司 equity administration",
            "侧重管理与估值，不强调跨文档法律证据驱动的异常定位",
        ],
    ],
    widths=[Inches(1.1), Inches(1.1), Inches(1.85), Inches(1.6), Inches(2.25)],
)

doc.add_heading("8. 重点竞品解读", level=1)
doc.add_heading("8.1 Carta", level=2)
add_bullets(
    doc,
    [
        "优势：品牌强、生态位强、cap table + 409A + liquidity + fund admin 一体化，适合被视为“单一真实数据源”。",
        "风险：如果客户已经在 Carta 中把数据治理做好，会削弱本产品的价值空间。",
        "机会：很多公司真正的问题在于历史材料并没有被完整、准确地录入系统；本产品可定位为“上架前审计层”和“交易前核查层”，补足现有系统的 blind spot。",
    ],
)

doc.add_heading("8.2 Pulley / Eqvista / Ledgy / Shareworks", level=2)
add_bullets(
    doc,
    [
        "这类产品都强调 cap table 清洁、合规、员工股权与实时报告，但主要假设系统内数据已经相对可信。",
        "本产品的差异点是把原始法律文件视为第一证据源，而不是默认系统台账为真。",
        "对于这类平台的现有用户，Cap Table Audit Copilot 可以作为“迁移前清洗工具”“融资前核对工具”或“审计补充层”切入。",
    ],
)

doc.add_heading("8.3 Kira / 法律尽调平台", level=2)
add_bullets(
    doc,
    [
        "Kira 等平台证明了律所愿意为结构化文档抽取和尽调效率买单。",
        "但它们更像“合同智能提取引擎”，不是专门解决股权事件重建和 cap table reconciliation 的产品。",
        "因此本产品无需与之争夺通用合同审阅市场，而应把自己定义为融资场景下的 vertical intelligence layer。",
    ],
)

doc.add_heading("9. 差异化定位", level=1)
add_bullets(
    doc,
    [
        "从“文档问答”升级为“交易智能与股权审计”。",
        "从“字段提取”升级为“事件流重建”。",
        "从“发现问题”升级为“定位证据 + 解释原因 + 标记缺失信息”。",
        "从“cap table 管理”升级为“cap table 可信度验证”。",
    ],
)
doc.add_paragraph(
    "一句话定位建议：AI-powered cap table reconciliation and equity audit copilot for startup financings。"
)

doc.add_heading("10. 市场规模判断（区间估算）", level=1)
doc.add_paragraph(
    "由于公开市场通常按 LegalTech、CLM、cap table management 或 equity management 统计，而非按“cap table 审计 copilot”统计，因此以下采用审慎估算。"
)
add_bullets(
    doc,
    [
        "Top-down：LegalTech 是一个数百亿美元市场，且 analytics、CLM、AI review 都在增长。本产品理论上属于高增长细分交叉带，而非边缘需求。",
        "SAM 假设：若优先覆盖北美与英国等创业融资活跃市场，目标客户以高频处理创业融资/尽调的律所团队、成长型公司法务/finance ops 为主，可服务市场大概率在“数亿美元级别”的软件与专业工作流预算带内。",
        "SOM 假设：hackathon 后 2-3 年内，如果以数十家律所团队和数十家高成长企业为目标，年客单价落在 6,000-30,000 美元区间，早期可验证收入空间约在百万美元级 ARR；若叠加按项目收费或按文档包收费，可进一步提升进入门槛较低的试点转化率。",
        "建议不要在对外 pitch 中过度强调极大市场，而应强调“高价值、强痛点、交易驱动、愿为效率和风险控制付费”的切口质量。",
    ],
)

doc.add_heading("11. 商业模式建议", level=1)
add_table(
    doc,
    ["模式", "适用对象", "优点", "风险", "建议"],
    [
        ["按席位订阅", "律所团队、法务团队", "收入稳定，利于留存", "早期门槛较高", "适合作为标准版"],
        ["按项目 / 数据室收费", "尽调项目、融资项目", "与用户价值强绑定，易试点", "收入波动较大", "适合作为切入方案"],
        ["基础订阅 + 用量计费", "高频团队", "兼顾扩张与成本回收", "定价设计复杂", "中期可采用"],
        ["审计报告专业版", "高风险交易、外部律师交付", "提升 ARPU", "需要更强模板和权限能力", "后续增值包"],
    ],
    widths=[Inches(1.35), Inches(1.45), Inches(1.55), Inches(1.45), Inches(1.1)],
)

doc.add_heading("12. Go-to-Market 建议", level=1)
add_numbered(
    doc,
    [
        "先卖“第一轮 review 加速器”，不要先卖“自动给法律结论”。",
        "用故意埋错的样例数据包演示三种异常：漏转换、缺批准、股份/比例不一致。",
        "优先进入愿意尝试新工具的 boutique / mid-size startup law firms，再复制到 in-house legal 与 finance ops。",
        "产品化包装上突出 evidence-linked audit report，而不是大而全的 AI assistant。",
        "销售策略上先以项目型试点切入，确认 ROI 后再推动年度订阅。",
    ],
)

doc.add_heading("13. 主要风险与应对", level=1)
add_table(
    doc,
    ["风险", "说明", "应对建议"],
    [
        [
            "规则复杂度失控",
            "SAFE / note / MFN / pro rata / multiple closing 很容易让规则爆炸",
            "MVP 严格限定 Delaware C-Corp 早期融资与常见证券类型",
        ],
        [
            "抽取准确率不稳定",
            "扫描件和版本冲突会直接影响结果可信度",
            "加入人工确认节点、置信度提示与缺失信息模型",
        ],
        [
            "客户担心法律责任",
            "若产品声称直接给出法律定论，风险过高",
            "定位为 audit copilot，输出“发现、原因、证据、待确认事项”",
        ],
        [
            "被平台型产品吸收",
            "Carta 等可能逐渐加入审计与 AI 功能",
            "抢先建立纵深工作流、证据链体验与律师场景口碑",
        ],
        [
            "采购周期偏长",
            "律所和法务 SaaS 采购常受安全与合规约束",
            "用项目化试点和私有化 / 权限控制能力降低进入阻力",
        ],
    ],
    widths=[Inches(1.2), Inches(2.4), Inches(2.8)],
)

doc.add_heading("14. 结论", level=1)
add_bullets(
    doc,
    [
        "这是一个较窄但质量很高的市场机会，核心价值不是“让律师少读一点字”，而是“让律师更快找到真正会影响交易的问题”。",
        "现有市场证明了两件事：第一，股权与 cap table 管理本身有稳定预算；第二，AI 驱动的尽调与合同分析已经被律所接受。",
        "真正的空白在两者之间：基于法律证据链的股权事件重建、理论 cap table 生成与异常解释。",
        "因此，Cap Table Audit Copilot 最适合以 Delaware startup financing 的高置信度审计工作流切入，并通过证据可追溯性、异常优先级、人与 AI 协作体验建立差异化。",
    ],
)

doc.add_heading("15. 建议的对外表述", level=1)
add_bullets(
    doc,
    [
        "We reconstruct what the cap table should be from legal documents, then compare it against what the company says it is.",
        "We do not just summarize documents; we detect inconsistencies, missing conversions, and missing approvals with evidence.",
        "We help startup financing lawyers and in-house teams find what humans are most likely to miss under time pressure.",
    ],
)

doc.add_heading("附录 A：研究来源", level=1)
sources = [
    (
        "Grand View Research",
        "Legal Technology Market Size, Share | Industry Report, 2033",
        "https://www.grandviewresearch.com/industry-analysis/legal-technology-market-report",
    ),
    (
        "American Bar Association",
        "How Lawyers Are Using AI to Draft Better Contracts Faster",
        "https://www.americanbar.org/groups/law_practice/resources/law-practice-today/2025/october-2025/lawyers-using-ai-to-draft-contracts/",
    ),
    (
        "Carta",
        "Equity Management Solutions - Equity Plans, Cap Tables & 409A / related product pages",
        "https://carta.com/landing/",
    ),
    (
        "Pulley",
        "Cap table and equity management / comparison and product materials",
        "https://new.pulley.com/",
    ),
    (
        "Ledgy",
        "Cap table management and compliance product pages",
        "https://ledgy.com/cap-table/",
    ),
    (
        "Morgan Stanley Shareworks",
        "Cap Table Management / Private Market Editions",
        "https://www.shareworks.com/product/cap-table-management-software/",
    ),
    (
        "Litera Kira",
        "Contract review and due diligence materials",
        "https://www.litera.com/blog/kira-remains-undisputed-leader-contract-review-now-supercharged-genai/",
    ),
    (
        "Eqvista",
        "Cap table and 409A related materials",
        "https://eqvista.com/free-cap-table-management-with-409a-valuation/",
    ),
]
for name, title_text, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{name} — {title_text} — {url}")

footer_section = doc.sections[0]
footer = footer_section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Cap Table Audit Copilot 市场调研报告")

doc.save(output_path)
print(output_path)
