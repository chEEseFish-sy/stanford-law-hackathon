#!/usr/bin/env python3
"""Process Series A DOCX files into structured JSON.

The script reads DOCX files from data_process, extracts local candidates, and
optionally calls Gemini to produce document-level structured data.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree


DEFAULT_INPUT_DIR = Path("data_process")
DEFAULT_OUTPUT_DIR = Path("processed_data")
DEFAULT_MODEL = "gemini-3-flash-preview"
MAX_CHUNK_CHARS = 18000
DEFAULT_TASKS = ["metadata", "dates", "financing", "rights", "risks"]


TASK_KEYWORDS = {
    "metadata": [
        "agreement",
        "company",
        "investor",
        "purchaser",
        "stockholder",
        "delaware",
    ],
    "dates": [
        "date",
        "deadline",
        "closing",
        "milestone",
        "effective",
        "initial closing",
        "additional closing",
        "tranche",
        "[●]",
    ],
    "financing": [
        "purchase price",
        "shares",
        "series a",
        "preferred stock",
        "safe",
        "cash purchaser",
        "valuation",
        "conversion",
        "liquidation",
        "tranche",
    ],
    "rights": [
        "information rights",
        "registration rights",
        "right of first refusal",
        "co-sale",
        "voting",
        "drag-along",
        "protective provisions",
        "preemptive rights",
        "observer",
        "transfer",
    ],
    "risks": [
        "default",
        "termination",
        "condition",
        "covenant",
        "breach",
        "waiver",
        "consent",
        "shall not",
        "unless",
        "provided that",
    ],
}


DATE_RE = re.compile(
    r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?)\.?\s+\d{1,2},\s+\d{4}\b"
    r"|\b\d{1,2}/\d{1,2}/\d{2,4}\b"
    r"|\b\d{4}-\d{1,2}-\d{1,2}\b"
    r"|\[[^\]]*●[^\]]*\]\s*,?\s*\d{4}"
    r"|\[[^\]]*●[^\]]*\]",
    re.IGNORECASE,
)
MONEY_RE = re.compile(
    r"(?:US\$|\$)\s?\d[\d,]*(?:\.\d+)?(?:\s?(?:million|billion|M|B|k))?"
    r"|\b\d[\d,]*(?:\.\d+)?\s?(?:dollars|USD)\b",
    re.IGNORECASE,
)
PERCENT_RE = re.compile(r"\b\d+(?:\.\d+)?\s?%")
SHARE_RE = re.compile(
    r"\b\d[\d,]*(?:\.\d+)?\s+(?:shares?|Shares?|Preferred Stock|Common Stock)\b"
)


@dataclass(frozen=True)
class TextBlock:
    paragraph_id: int
    text: str
    block_type: str = "paragraph"


@dataclass(frozen=True)
class Chunk:
    chunk_id: int
    start_paragraph_id: int
    end_paragraph_id: int
    text: str


def safe_stem(path: Path) -> str:
    stem = path.stem.strip().lower()
    stem = re.sub(r"[^a-z0-9]+", "_", stem)
    return stem.strip("_") or "document"


def load_dotenv_if_available() -> None:
    try:
        from dotenv import load_dotenv
    except ImportError:
        return
    load_dotenv()


def parse_docx(path: Path) -> list[TextBlock]:
    try:
        return parse_docx_with_python_docx(path)
    except ImportError:
        return parse_docx_with_zip(path)


def parse_docx_with_python_docx(path: Path) -> list[TextBlock]:
    from docx import Document

    doc = Document(path)
    blocks: list[TextBlock] = []
    paragraph_id = 0

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            blocks.append(TextBlock(paragraph_id=paragraph_id, text=text))
            paragraph_id += 1

    for table_index, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([clean_text(cell.text) for cell in row.cells])
        text = json.dumps(rows, ensure_ascii=False)
        if text and rows:
            blocks.append(
                TextBlock(
                    paragraph_id=paragraph_id,
                    text=f"TABLE {table_index}: {text}",
                    block_type="table",
                )
            )
            paragraph_id += 1

    return blocks


def parse_docx_with_zip(path: Path) -> list[TextBlock]:
    """Fallback parser used when python-docx is not installed."""

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    blocks: list[TextBlock] = []
    paragraph_id = 0

    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))

    body = root.find(".//w:body", ns)
    if body is None:
        return []

    for child in list(body):
        if child.tag.endswith("}p"):
            text = clean_text("".join(t.text or "" for t in child.findall(".//w:t", ns)))
            if text:
                blocks.append(TextBlock(paragraph_id=paragraph_id, text=text))
                paragraph_id += 1
        elif child.tag.endswith("}tbl"):
            rows: list[list[str]] = []
            for row in child.findall(".//w:tr", ns):
                cells: list[str] = []
                for cell in row.findall("./w:tc", ns):
                    cells.append(clean_text("".join(t.text or "" for t in cell.findall(".//w:t", ns))))
                if any(cells):
                    rows.append(cells)
            if rows:
                blocks.append(
                    TextBlock(
                        paragraph_id=paragraph_id,
                        text=f"TABLE: {json.dumps(rows, ensure_ascii=False)}",
                        block_type="table",
                    )
                )
                paragraph_id += 1

    return blocks


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def make_chunks(blocks: list[TextBlock], max_chars: int = MAX_CHUNK_CHARS) -> list[Chunk]:
    chunks: list[Chunk] = []
    current: list[str] = []
    start_id: int | None = None
    end_id: int | None = None
    current_len = 0

    for block in blocks:
        line = f"[paragraph_id={block.paragraph_id} type={block.block_type}] {block.text}"
        line_len = len(line) + 1

        if current and current_len + line_len > max_chars:
            chunks.append(
                Chunk(
                    chunk_id=len(chunks),
                    start_paragraph_id=start_id if start_id is not None else 0,
                    end_paragraph_id=end_id if end_id is not None else 0,
                    text="\n".join(current),
                )
            )
            current = []
            start_id = None
            end_id = None
            current_len = 0

        if start_id is None:
            start_id = block.paragraph_id
        end_id = block.paragraph_id
        current.append(line)
        current_len += line_len

    if current:
        chunks.append(
            Chunk(
                chunk_id=len(chunks),
                start_paragraph_id=start_id if start_id is not None else 0,
                end_paragraph_id=end_id if end_id is not None else 0,
                text="\n".join(current),
            )
        )

    return chunks


def extract_raw_candidates(blocks: list[TextBlock]) -> dict[str, list[dict[str, Any]]]:
    return {
        "dates": extract_matches(blocks, DATE_RE, normalize_dates=True),
        "money": extract_matches(blocks, MONEY_RE),
        "percentages": extract_matches(blocks, PERCENT_RE),
        "share_counts": extract_matches(blocks, SHARE_RE),
        "section_titles": extract_section_titles(blocks),
    }


def build_candidate_groups(blocks: list[TextBlock]) -> dict[str, list[dict[str, Any]]]:
    groups: dict[str, list[dict[str, Any]]] = {}
    for task, keywords in TASK_KEYWORDS.items():
        groups[task] = select_candidate_paragraphs(blocks, keywords)
    return groups


def select_candidate_paragraphs(
    blocks: list[TextBlock],
    keywords: list[str],
    max_items: int = 80,
    context_window: int = 1,
) -> list[dict[str, Any]]:
    selected_ids: set[int] = set()
    by_id = {block.paragraph_id: block for block in blocks}

    for block in blocks:
        text_lower = block.text.lower()
        if any(keyword.lower() in text_lower for keyword in keywords):
            for paragraph_id in range(
                max(0, block.paragraph_id - context_window),
                block.paragraph_id + context_window + 1,
            ):
                if paragraph_id in by_id:
                    selected_ids.add(paragraph_id)

    return [
        {
            "paragraph_id": paragraph_id,
            "type": by_id[paragraph_id].block_type,
            "text": by_id[paragraph_id].text,
        }
        for paragraph_id in sorted(selected_ids)[:max_items]
    ]


def extract_matches(
    blocks: list[TextBlock], pattern: re.Pattern[str], normalize_dates: bool = False
) -> list[dict[str, Any]]:
    seen: set[tuple[str, int]] = set()
    matches: list[dict[str, Any]] = []

    for block in blocks:
        for match in pattern.finditer(block.text):
            value = clean_text(match.group(0))
            key = (value.lower(), block.paragraph_id)
            if key in seen:
                continue
            seen.add(key)
            item: dict[str, Any] = {
                "value": value,
                "paragraph_id": block.paragraph_id,
                "source_text": trim_evidence(block.text),
            }
            if normalize_dates:
                item["date_iso"] = normalize_date(value)
            matches.append(item)

    return matches[:200]


def extract_section_titles(blocks: list[TextBlock]) -> list[dict[str, Any]]:
    titles: list[dict[str, Any]] = []
    for block in blocks:
        text = block.text
        looks_like_title = (
            len(text) <= 120
            and not text.endswith(";")
            and (
                text.endswith(".")
                or text.isupper()
                or re.match(r"^\d+(?:\.\d+)*\.?\s+[A-Z]", text)
            )
        )
        if looks_like_title:
            titles.append(
                {
                    "value": text,
                    "paragraph_id": block.paragraph_id,
                    "source_text": trim_evidence(text),
                }
            )
    return titles[:100]


def normalize_date(value: str) -> str | None:
    if "●" in value:
        return None
    try:
        import dateparser
    except ImportError:
        return None

    parsed = dateparser.parse(value, settings={"PREFER_DAY_OF_MONTH": "first"})
    if parsed is None:
        return None
    return parsed.date().isoformat()


def trim_evidence(text: str, limit: int = 600) -> str:
    text = clean_text(text)
    if len(text) <= limit:
        return text
    return text[: limit - 3].rstrip() + "..."


def empty_document_result(
    source_file: str,
    model: str,
    chunk_count: int,
    raw_candidates: dict[str, Any],
    status: str = "processed",
) -> dict[str, Any]:
    return {
        "source_file": source_file,
        "document_title": "",
        "document_type": "",
        "parties": [],
        "dates": [],
        "financing_terms": {
            "round": "",
            "security_classes": [],
            "purchase_price": [],
            "share_amounts": [],
            "valuation_or_cap": [],
            "closing_terms": [],
            "milestones": [],
        },
        "investor_rights": [],
        "company_obligations": [],
        "transfer_or_voting_restrictions": [],
        "key_conditions": [],
        "risks_or_review_items": [],
        "raw_candidates": raw_candidates,
        "confidence": 0.0,
        "processing_metadata": {
            "model": model,
            "processed_at": datetime.now(timezone.utc).isoformat(),
            "chunk_count": chunk_count,
            "status": status,
        },
    }


class GeminiExtractor:
    def __init__(self, api_key: str, model: str) -> None:
        try:
            from google import genai
            from google.genai import types
        except ImportError as exc:
            raise RuntimeError(
                "Missing Gemini SDK. Install dependencies with: "
                "python3 -m pip install -r requirements.txt"
            ) from exc

        self.client = genai.Client(api_key=api_key)
        self.model = model
        self.types = types

    def generate_json(self, prompt: str) -> dict[str, Any]:
        response = self.client.models.generate_content(
            model=self.model,
            contents=prompt,
            config=self.types.GenerateContentConfig(
                temperature=0,
                response_mime_type="application/json",
            ),
        )
        text = getattr(response, "text", None) or ""
        return parse_json_response(text)


def parse_json_response(text: str) -> dict[str, Any]:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise
        parsed = json.loads(text[start : end + 1])
    if not isinstance(parsed, dict):
        raise ValueError("Gemini response must be a JSON object")
    return parsed


def build_task_prompt(
    source_file: str,
    task: str,
    candidate_paragraphs: list[dict[str, Any]],
    raw_candidates: dict[str, Any],
) -> str:
    schema = task_schema(task)
    prompt_candidates = {
        key: value[:40] if isinstance(value, list) else value for key, value in raw_candidates.items()
    }
    return f"""
You extract one focused category of structured data from a venture financing legal agreement.
Return only valid JSON. Do not include Markdown.
Do not invent facts that are not supported by the provided paragraph text.
Every important extracted item must include evidence with paragraph_id and source_text.
If a date is a placeholder such as [●], keep date_original and set date_iso to null.

Source file: {source_file}
Task: {task}

Required JSON shape:
{json.dumps(schema, ensure_ascii=False, indent=2)}

Local regex candidates, for reference only:
{json.dumps(prompt_candidates, ensure_ascii=False, indent=2)}

Candidate paragraphs:
{json.dumps(candidate_paragraphs, ensure_ascii=False, indent=2)}
""".strip()


def task_schema(task: str) -> dict[str, Any]:
    schemas: dict[str, dict[str, Any]] = {
        "metadata": {
            "document_title": "",
            "document_type": "",
            "parties": [
                {
                    "name": "",
                    "role": "",
                    "evidence": [{"paragraph_id": 0, "source_text": ""}],
                }
            ],
            "confidence": 0.0,
        },
        "dates": {
            "dates": [
                {
                    "date_original": "",
                    "date_iso": None,
                    "meaning": "",
                    "evidence": [{"paragraph_id": 0, "source_text": ""}],
                }
            ],
            "financing_terms": {"closing_terms": [], "milestones": []},
            "confidence": 0.0,
        },
        "financing": {
            "financing_terms": {
                "round": "",
                "security_classes": [],
                "purchase_price": [],
                "share_amounts": [],
                "valuation_or_cap": [],
                "closing_terms": [],
                "milestones": [],
            },
            "confidence": 0.0,
        },
        "rights": {
            "investor_rights": [],
            "company_obligations": [],
            "transfer_or_voting_restrictions": [],
            "key_conditions": [],
            "confidence": 0.0,
        },
        "risks": {
            "risks_or_review_items": [],
            "key_conditions": [],
            "confidence": 0.0,
        },
    }
    return schemas[task]


def deterministic_merge(
    source_file: str,
    first_title: str,
    model: str,
    chunk_count: int,
    raw_candidates: dict[str, Any],
    chunk_results: list[dict[str, Any]],
    status: str = "processed",
) -> dict[str, Any]:
    result = empty_document_result(source_file, model, chunk_count, raw_candidates, status=status)
    result["document_title"] = first_title

    for chunk_result in chunk_results:
        if not isinstance(chunk_result, dict):
            continue
        for key in ["document_type"]:
            if not result[key] and isinstance(chunk_result.get(key), str):
                result[key] = chunk_result[key]
        if not result["document_title"] and isinstance(chunk_result.get("document_title"), str):
            result["document_title"] = chunk_result["document_title"]
        result["parties"] = merge_list(result["parties"], chunk_result.get("parties"))
        result["dates"] = merge_list(result["dates"], chunk_result.get("dates"))
        for field in [
            "investor_rights",
            "company_obligations",
            "transfer_or_voting_restrictions",
            "key_conditions",
            "risks_or_review_items",
        ]:
            result[field] = merge_list(result[field], chunk_result.get(field))
        financing = chunk_result.get("financing_terms")
        if isinstance(financing, dict):
            for key, value in financing.items():
                if key == "round" and not result["financing_terms"]["round"] and isinstance(value, str):
                    result["financing_terms"]["round"] = value
                elif key in result["financing_terms"]:
                    result["financing_terms"][key] = merge_list(result["financing_terms"][key], value)

    confidences = [
        item.get("confidence")
        for item in chunk_results
        if isinstance(item, dict) and isinstance(item.get("confidence"), (int, float))
    ]
    if confidences:
        result["confidence"] = round(sum(confidences) / len(confidences), 2)
    return result


def merge_list(existing: list[Any], incoming: Any) -> list[Any]:
    if not isinstance(incoming, list):
        return existing
    seen = {stable_key(item) for item in existing}
    for item in incoming:
        key = stable_key(item)
        if key not in seen:
            existing.append(item)
            seen.add(key)
    return existing


def stable_key(value: Any) -> str:
    if isinstance(value, dict):
        compact = {
            key: value.get(key)
            for key in sorted(value)
            if key not in {"evidence", "source_text"}
        }
        return json.dumps(compact, ensure_ascii=False, sort_keys=True)
    return json.dumps(value, ensure_ascii=False, sort_keys=True)


def process_file(
    path: Path,
    output_dir: Path,
    use_llm: bool,
    model: str,
    max_chars: int,
    tasks: list[str],
) -> dict[str, Any]:
    blocks = parse_docx(path)
    chunks = make_chunks(blocks, max_chars=max_chars)
    raw_candidates = extract_raw_candidates(blocks)
    candidate_groups = build_candidate_groups(blocks)
    first_title = blocks[0].text if blocks else ""
    source_file = path.name

    parsed_payload = {
        "source_file": source_file,
        "document_title": first_title,
        "paragraph_count": len(blocks),
        "character_count": sum(len(block.text) for block in blocks),
        "chunks": [
            {
                "chunk_id": chunk.chunk_id,
                "start_paragraph_id": chunk.start_paragraph_id,
                "end_paragraph_id": chunk.end_paragraph_id,
                "character_count": len(chunk.text),
            }
            for chunk in chunks
        ],
        "paragraphs": [
            {
                "paragraph_id": block.paragraph_id,
                "type": block.block_type,
                "text": block.text,
            }
            for block in blocks
        ],
    }
    candidates_payload = {
        "source_file": source_file,
        "document_title": first_title,
        "raw_candidates": raw_candidates,
        "candidate_groups": candidate_groups,
        "processing_metadata": {
            "processed_at": datetime.now(timezone.utc).isoformat(),
            "status": "candidates_ready",
            "paragraph_count": len(blocks),
            "character_count": sum(len(block.text) for block in blocks),
        },
    }

    write_json(output_dir / "parsed" / f"{safe_stem(path)}.parsed.json", parsed_payload)
    write_json(output_dir / "candidates" / f"{safe_stem(path)}.candidates.json", candidates_payload)

    if not use_llm:
        result = empty_document_result(
            source_file=source_file,
            model="no-llm",
            chunk_count=len(chunks),
            raw_candidates=raw_candidates,
            status="candidates_ready",
        )
        result["document_title"] = first_title
        result["processing_metadata"]["paragraph_count"] = len(blocks)
        result["processing_metadata"]["character_count"] = sum(len(block.text) for block in blocks)
        result["candidate_group_counts"] = {
            task: len(items) for task, items in candidate_groups.items()
        }
    else:
        load_dotenv_if_available()
        api_key = os.getenv("LLM_API_KEY")
        if not api_key:
            raise RuntimeError("LLM_API_KEY is missing. Add it to .env or export it.")
        extractor = GeminiExtractor(api_key=api_key, model=model)

        task_results: list[dict[str, Any]] = []
        errors: list[dict[str, Any]] = []
        for task in tasks:
            paragraphs = candidate_groups.get(task, [])
            if not paragraphs:
                continue
            print(f"  LLM task {task}: {len(paragraphs)} candidate paragraph(s)", flush=True)
            try:
                task_result = extractor.generate_json(
                    build_task_prompt(source_file, task, paragraphs, raw_candidates)
                )
                task_result["_task"] = task
                task_results.append(task_result)
            except Exception as exc:  # noqa: BLE001 - keep batch processing alive.
                errors.append({"task": task, "error": str(exc)})

        if task_results:
            result = deterministic_merge(
                source_file,
                first_title,
                model,
                len(chunks),
                raw_candidates,
                task_results,
                status="processed_with_task_llm",
            )
        else:
            result = empty_document_result(
                source_file,
                model,
                len(chunks),
                raw_candidates,
                status="failed",
            )

        if errors:
            result["processing_metadata"]["errors"] = errors
        result["processing_metadata"]["paragraph_count"] = len(blocks)
        result["processing_metadata"]["character_count"] = sum(len(block.text) for block in blocks)
        result["processing_metadata"]["tasks"] = tasks
        result["candidate_group_counts"] = {
            task: len(items) for task, items in candidate_groups.items()
        }

    output_path = output_dir / f"{safe_stem(path)}.json"
    write_json(output_path, result)
    return result


def normalize_result_shape(
    result: dict[str, Any],
    source_file: str,
    model: str,
    chunk_count: int,
    raw_candidates: dict[str, Any],
) -> dict[str, Any]:
    base = empty_document_result(source_file, model, chunk_count, raw_candidates)
    for key, value in result.items():
        if key == "processing_metadata" and isinstance(value, dict):
            base[key].update(value)
        elif key == "financing_terms" and isinstance(value, dict):
            base[key].update({k: v for k, v in value.items() if k in base[key]})
        elif key in base:
            base[key] = value
    base["source_file"] = source_file
    base["raw_candidates"] = raw_candidates
    base["processing_metadata"]["model"] = model
    base["processing_metadata"]["processed_at"] = datetime.now(timezone.utc).isoformat()
    base["processing_metadata"]["chunk_count"] = chunk_count
    base["processing_metadata"].setdefault("status", "processed")
    return base


def build_index(results: list[dict[str, Any]]) -> dict[str, Any]:
    documents = []
    for result in results:
        parties = result.get("parties") if isinstance(result.get("parties"), list) else []
        dates = result.get("dates") if isinstance(result.get("dates"), list) else []
        metadata = result.get("processing_metadata") if isinstance(result.get("processing_metadata"), dict) else {}
        documents.append(
            {
                "source_file": result.get("source_file", ""),
                "document_title": result.get("document_title", ""),
                "document_type": result.get("document_type", ""),
                "key_dates": dates[:10],
                "main_parties": parties[:10],
                "confidence": result.get("confidence", 0.0),
                "status": metadata.get("status", "processed"),
                "output_file": f"{safe_stem(Path(str(result.get('source_file', 'document'))))}.json",
            }
        )
    return {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "document_count": len(documents),
        "documents": documents,
    }


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def iter_input_files(input_dir: Path, single_file: Path | None) -> list[Path]:
    if single_file is not None:
        return [single_file]
    return sorted(path for path in input_dir.glob("*.docx") if not path.name.startswith("~$"))


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract structured data from DOCX financing documents.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT_DIR, help="Input directory containing DOCX files.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT_DIR, help="Output directory for JSON files.")
    parser.add_argument("--file", type=Path, default=None, help="Process one DOCX file instead of the input directory.")
    parser.add_argument("--llm", action="store_true", help="Call Gemini on task-specific candidate paragraphs.")
    parser.add_argument(
        "--no-llm",
        action="store_true",
        help="Deprecated alias for the default behavior. No Gemini calls are made unless --llm is set.",
    )
    parser.add_argument(
        "--task",
        choices=[*DEFAULT_TASKS, "all"],
        action="append",
        default=None,
        help="LLM task to run. Repeat for multiple tasks. Defaults to all tasks when --llm is set.",
    )
    parser.add_argument("--max-chars", type=int, default=MAX_CHUNK_CHARS, help="Maximum characters per LLM chunk.")
    return parser.parse_args(argv)


def main(argv: list[str]) -> int:
    args = parse_args(argv)
    global MAX_CHUNK_CHARS
    MAX_CHUNK_CHARS = args.max_chars

    load_dotenv_if_available()
    model = os.getenv("LLM_MODEL_NAME", DEFAULT_MODEL)
    tasks = DEFAULT_TASKS if not args.task or "all" in args.task else args.task
    files = iter_input_files(args.input, args.file)
    if not files:
        print(f"No DOCX files found in {args.input}", file=sys.stderr)
        return 1

    args.output.mkdir(parents=True, exist_ok=True)
    results: list[dict[str, Any]] = []
    failures: list[dict[str, str]] = []

    for path in files:
        print(f"Processing {path}...")
        try:
            results.append(process_file(path, args.output, args.llm, model, args.max_chars, tasks))
        except Exception as exc:  # noqa: BLE001
            failures.append({"source_file": str(path), "error": str(exc)})
            print(f"Failed {path}: {exc}", file=sys.stderr)

    index = build_index(results)
    if failures:
        index["failures"] = failures
    write_json(args.output / "index.json", index)
    print(f"Wrote {len(results)} document result(s) to {args.output}")
    failed_results = [
        result
        for result in results
        if isinstance(result.get("processing_metadata"), dict)
        and result["processing_metadata"].get("status") == "failed"
    ]
    if failures:
        print(f"{len(failures)} file(s) failed. See processed_data/index.json.", file=sys.stderr)
        return 2
    if failed_results:
        print(
            f"{len(failed_results)} document(s) had failed processing status. "
            f"See {args.output / 'index.json'}.",
            file=sys.stderr,
        )
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
